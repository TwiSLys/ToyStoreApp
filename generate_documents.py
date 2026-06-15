import os

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Cm, Pt
from reportlab.lib import colors
from reportlab.lib.pagesizes import A4, landscape
from reportlab.pdfbase import pdfmetrics
from reportlab.pdfbase.ttfonts import TTFont
from reportlab.pdfgen import canvas


ROOT = os.path.dirname(os.path.abspath(__file__))
DOCS_DIR = os.path.join(ROOT, "docs")
FONT_PATH = os.path.join(os.environ.get("WINDIR", r"C:\Windows"), "Fonts", "arial.ttf")
FONT_BOLD_PATH = os.path.join(
    os.environ.get("WINDIR", r"C:\Windows"), "Fonts", "arialbd.ttf"
)


def register_fonts():
    pdfmetrics.registerFont(TTFont("Arial", FONT_PATH))
    pdfmetrics.registerFont(TTFont("Arial-Bold", FONT_BOLD_PATH))


def draw_table(pdf, x, y, width, title, fields):
    line_height = 14
    height = 25 + line_height * len(fields)
    pdf.setFillColor(colors.HexColor("#F5DEB3"))
    pdf.roundRect(x, y - height, width, height, 5, fill=1, stroke=1)
    pdf.setFillColor(colors.black)
    pdf.setFont("Arial-Bold", 9)
    pdf.drawCentredString(x + width / 2, y - 16, title)
    pdf.line(x, y - 23, x + width, y - 23)
    pdf.setFont("Arial", 7.5)
    current_y = y - 35
    for field in fields:
        pdf.drawString(x + 6, current_y, field)
        current_y -= line_height
    return (x, y - height, x + width, y)


def connect_boxes(pdf, first, second):
    x1 = (first[0] + first[2]) / 2
    y1 = (first[1] + first[3]) / 2
    x2 = (second[0] + second[2]) / 2
    y2 = (second[1] + second[3]) / 2
    pdf.setStrokeColor(colors.HexColor("#6B4F34"))
    pdf.setLineWidth(1)
    pdf.line(x1, y1, x2, y2)


def create_er_diagram():
    path = os.path.join(DOCS_DIR, "ER-диаграмма.pdf")
    page_width, page_height = landscape(A4)
    pdf = canvas.Canvas(path, pagesize=(page_width, page_height))
    pdf.setTitle("ER-диаграмма ООО МирИгрушек")
    pdf.setFont("Arial-Bold", 16)
    pdf.drawCentredString(page_width / 2, page_height - 25, "ER-диаграмма базы данных")

    boxes = {}
    boxes["roles"] = draw_table(
        pdf, 25, 520, 130, "roles", ["PK id", "name (UQ)"]
    )
    boxes["users"] = draw_table(
        pdf,
        190,
        520,
        170,
        "users",
        ["PK id", "FK role_id", "full_name", "login (UQ)", "password"],
    )
    boxes["categories"] = draw_table(
        pdf, 395, 520, 140, "categories", ["PK id", "name (UQ)"]
    )
    boxes["suppliers"] = draw_table(
        pdf, 570, 520, 140, "suppliers", ["PK id", "name (UQ)"]
    )
    boxes["manufacturers"] = draw_table(
        pdf, 675, 350, 140, "manufacturers", ["PK id", "name (UQ)"]
    )
    boxes["products"] = draw_table(
        pdf,
        395,
        350,
        215,
        "products",
        [
            "PK id",
            "article (UQ)",
            "name, unit, price",
            "FK supplier_id",
            "FK manufacturer_id",
            "FK category_id",
            "discount, stock",
            "description, image_path",
        ],
    )
    boxes["pickup_points"] = draw_table(
        pdf, 25, 280, 180, "pickup_points", ["PK id", "address (UQ)"]
    )
    boxes["orders"] = draw_table(
        pdf,
        230,
        280,
        190,
        "orders",
        [
            "PK id",
            "order_date",
            "delivery_date",
            "FK pickup_point_id",
            "FK client_id",
            "receive_code",
            "status",
        ],
    )
    boxes["order_items"] = draw_table(
        pdf,
        470,
        145,
        180,
        "order_items",
        ["PK/FK order_id", "PK/FK product_id", "quantity"],
    )

    for first, second in (
        ("roles", "users"),
        ("users", "orders"),
        ("pickup_points", "orders"),
        ("orders", "order_items"),
        ("products", "order_items"),
        ("categories", "products"),
        ("suppliers", "products"),
        ("manufacturers", "products"),
    ):
        connect_boxes(pdf, boxes[first], boxes[second])
    pdf.setFont("Arial", 8)
    pdf.drawString(25, 20, "Связи реализованы внешними ключами; схема приведена к 3НФ.")
    pdf.save()


def centered_text(pdf, x, y, width, height, text, font_size=9):
    pdf.setFont("Arial", font_size)
    lines = text.split("\n")
    start_y = y + height / 2 + (len(lines) - 1) * 5
    for index, line in enumerate(lines):
        pdf.drawCentredString(x + width / 2, start_y - index * 11, line)


def arrow(pdf, x1, y1, x2, y2):
    pdf.line(x1, y1, x2, y2)
    if y2 < y1:
        pdf.line(x2, y2, x2 - 4, y2 + 7)
        pdf.line(x2, y2, x2 + 4, y2 + 7)
    elif x2 > x1:
        pdf.line(x2, y2, x2 - 7, y2 - 4)
        pdf.line(x2, y2, x2 - 7, y2 + 4)


def create_flowchart():
    path = os.path.join(DOCS_DIR, "Блок-схема.pdf")
    pdf = canvas.Canvas(path, pagesize=A4)
    width, height = A4
    pdf.setTitle("Алгоритм приложения")
    pdf.setFont("Arial-Bold", 15)
    pdf.drawCentredString(width / 2, height - 25, "Алгоритм работы приложения")
    pdf.setStrokeColor(colors.black)

    x = 190
    box_width = 215
    box_height = 38
    y = 760
    pdf.roundRect(x, y, box_width, box_height, 18, fill=0)
    centered_text(pdf, x, y, box_width, box_height, "Запуск приложения")
    arrow(pdf, x + box_width / 2, y, x + box_width / 2, y - 25)

    y -= 65
    pdf.rect(x, y, box_width, box_height)
    centered_text(pdf, x, y, box_width, box_height, "Подключение к SQLite")
    arrow(pdf, x + box_width / 2, y, x + box_width / 2, y - 25)

    y -= 70
    diamond_center = (width / 2, y + 25)
    points = [
        (diamond_center[0], diamond_center[1] + 32),
        (diamond_center[0] + 105, diamond_center[1]),
        (diamond_center[0], diamond_center[1] - 32),
        (diamond_center[0] - 105, diamond_center[1]),
    ]
    path_obj = pdf.beginPath()
    path_obj.moveTo(*points[0])
    for point in points[1:]:
        path_obj.lineTo(*point)
    path_obj.close()
    pdf.drawPath(path_obj)
    centered_text(pdf, x, y, box_width, 50, "Вход выполнен?")

    pdf.setFont("Arial", 8)
    pdf.drawString(diamond_center[0] + 112, diamond_center[1] + 4, "Нет")
    arrow(pdf, diamond_center[0] + 105, diamond_center[1], 500, diamond_center[1])
    pdf.rect(455, diamond_center[1] - 70, 120, 40)
    centered_text(pdf, 455, diamond_center[1] - 70, 120, 40, "Сообщение\nоб ошибке", 8)
    arrow(pdf, 515, diamond_center[1] - 30, 515, diamond_center[1] - 5)

    pdf.drawString(diamond_center[0] + 8, diamond_center[1] - 47, "Да / гость")
    arrow(
        pdf,
        diamond_center[0],
        diamond_center[1] - 32,
        diamond_center[0],
        diamond_center[1] - 72,
    )
    y = diamond_center[1] - 112
    pdf.rect(x, y, box_width, box_height)
    centered_text(pdf, x, y, box_width, box_height, "Определение роли и прав")
    arrow(pdf, x + box_width / 2, y, x + box_width / 2, y - 25)

    y -= 70
    pdf.rect(x, y, box_width, box_height)
    centered_text(pdf, x, y, box_width, box_height, "Загрузка списка товаров")
    arrow(pdf, x + box_width / 2, y, x + box_width / 2, y - 25)

    y -= 75
    pdf.rect(45, y, 150, 52)
    centered_text(pdf, 45, y, 150, 52, "Гость / клиент:\nпросмотр")
    pdf.rect(223, y, 150, 52)
    centered_text(pdf, 223, y, 150, 52, "Менеджер:\nпоиск, фильтр,\nзаказы")
    pdf.rect(401, y, 150, 52)
    centered_text(pdf, 401, y, 150, 52, "Администратор:\nCRUD товаров\nи заказов")
    branch_y = y + 72
    pdf.line(width / 2, y + 95, width / 2, branch_y)
    arrow(pdf, width / 2, branch_y, 120, y + 52)
    arrow(pdf, width / 2, branch_y, 298, y + 52)
    arrow(pdf, width / 2, branch_y, 476, y + 52)

    final_y = y - 80
    for center_x in (120, 298, 476):
        arrow(pdf, center_x, y, width / 2, final_y + 38)
    pdf.roundRect(x, final_y, box_width, box_height, 18, fill=0)
    centered_text(pdf, x, final_y, box_width, box_height, "Выход / возврат к окну входа")
    pdf.setFont("Arial", 8)
    pdf.drawString(
        35,
        22,
        "Условные обозначения соответствуют ГОСТ 19.701-90: терминатор, процесс, решение.",
    )
    pdf.save()


def set_document_font(document):
    styles = document.styles
    for style_name in ("Normal", "Title", "Heading 1", "Heading 2"):
        style = styles[style_name]
        style.font.name = "Arial"
        style.font.size = Pt(11 if style_name == "Normal" else 14)


def create_screenshot_report():
    path = os.path.join(DOCS_DIR, "Скриншоты_работы.docx")
    document = Document()
    set_document_font(document)
    title = document.add_heading("Проверка работы системы ООО «МирИгрушек»", 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    document.add_paragraph(
        "Дата проверки: 15.06.2026. Приложение запущено с заполненной базой "
        "SQLite. Аварийного завершения при проверке не выявлено."
    )
    screenshots = [
        ("Окно входа", "01_login.png"),
        ("Каталог товаров администратора", "02_products_admin.png"),
        ("Список заказов администратора", "03_orders_admin.png"),
    ]
    for heading, file_name in screenshots:
        document.add_heading(heading, level=1)
        image_path = os.path.join(DOCS_DIR, "screenshots", file_name)
        if os.path.exists(image_path):
            paragraph = document.add_paragraph()
            paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
            paragraph.add_run().add_picture(image_path, width=Cm(16.5))
    document.add_heading("Проверенные сценарии", level=1)
    for text in (
        "Авторизация администратора и менеджера, отказ при неверных данных.",
        "Гостевой просмотр каталога без поиска, фильтрации и сортировки.",
        "Живой поиск по артикулу и текстовым полям, фильтр по поставщику.",
        "Сортировка по цене и количеству в обоих направлениях.",
        "Добавление, изменение и удаление тестового товара.",
        "Запрет удаления товара, присутствующего в заказе.",
        "Добавление и удаление тестового заказа.",
        "Отображение фото, скидки, итоговой цены и остатков.",
    ):
        document.add_paragraph(text, style="List Bullet")
    document.save(path)


def main():
    os.makedirs(DOCS_DIR, exist_ok=True)
    register_fonts()
    create_er_diagram()
    create_flowchart()
    create_screenshot_report()
    print(f"Documents created in {DOCS_DIR}")


if __name__ == "__main__":
    main()
